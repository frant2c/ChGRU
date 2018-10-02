import sys
import math
from itertools import groupby

from PyQt5.QtCore import pyqtSlot
from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QPushButton, QLabel, QLineEdit, QGridLayout, QTextEdit, QFileDialog, QMessageBox
from PyQt5.QtGui import QIntValidator, QDoubleValidator

from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class Window1(QWidget):
    def __init__(self):
        super(Window1, self).__init__()

        authors_name = QLabel("Исполнитель")
        well = QLabel('Номер скважины')
        locat = QLabel('Поднятие')
        altust = QLabel('Альтитуда устья')
        altrot = QLabel('Альтитуда ротора')
        drillstart = QLabel('Начало бурения')
        drillend = QLabel('Окончание бурения')
        depth_proj = QLabel('Глубина проектная')
        depth_fact = QLabel('Глубина фактическая')
        horizont_proj = QLabel('Горизонт проектный')
        horizont_fact = QLabel('Горизонт фактический')
        well_place = QLabel('Местоположение скважины:')
        diametr = QLabel('Диаметр ствола скважины:')
        lbl_1 = QLabel('Количество рейсов')
        lbl_2 = QLabel('Количество слоев')
        customer = QLabel('Заказчик')
        shell = QLabel('Снаряд')
        cat_well = QLabel('Категория скважины')

        self.authors_nameEdit = QLineEdit()
        self.wellEdit = QLineEdit()
        self.locatEdit = QLineEdit()
        self.altustEdit = QLineEdit()
        self.altrotEdit = QLineEdit()
        self.drillstartEdit = QLineEdit()
        self.drillendEdit = QLineEdit()
        self.depth_projEdit = QLineEdit()
        self.depth_factEdit = QLineEdit()
        self.horizont_projEdit = QLineEdit()
        self.horizont_factEdit = QLineEdit()
        self.well_placeEdit = QLineEdit()
        self.diametrEdit = QLineEdit()
        self.reysis = QLineEdit()
        self.reysis.setValidator(QIntValidator(1, 100))
        self.intervals = QLineEdit()
        self.intervals.setValidator(QIntValidator(1, 100))
        self.customerEdit = QLineEdit()
        self.shellEdit = QLineEdit()
        self.cat_wellEdit = QLineEdit()
        self.button = QPushButton(self)
        self.button.setText('Ok')
        self.button.clicked.connect(self.extraction_1)

        grid = QGridLayout()
        grid.addWidget(authors_name, 0, 0)
        grid.addWidget(self.authors_nameEdit, 0, 1)
        grid.addWidget(well, 1, 0)
        grid.addWidget(self.wellEdit, 1, 1)
        grid.addWidget(locat, 2, 0)
        grid.addWidget(self.locatEdit, 2, 1)
        grid.addWidget(altust, 3, 0)
        grid.addWidget(self.altustEdit, 3, 1)
        grid.addWidget(altrot, 4, 0)
        grid.addWidget(self.altrotEdit, 4, 1)
        grid.addWidget(drillstart, 5, 0)
        grid.addWidget(self.drillstartEdit, 5, 1)
        grid.addWidget(drillend, 6, 0)
        grid.addWidget(self.drillendEdit, 6, 1)
        grid.addWidget(depth_proj, 7, 0)
        grid.addWidget(self.depth_projEdit, 7, 1)
        grid.addWidget(depth_fact, 8, 0)
        grid.addWidget(self.depth_factEdit, 8, 1)
        grid.addWidget(horizont_proj, 9, 0)
        grid.addWidget(self.horizont_projEdit, 9, 1)
        grid.addWidget(horizont_fact, 10, 0)
        grid.addWidget(self.horizont_factEdit, 10, 1)
        grid.addWidget(well_place, 11, 0)
        grid.addWidget(self.well_placeEdit, 11, 1)
        grid.addWidget(diametr, 12, 0)
        grid.addWidget(self.diametrEdit, 12, 1)
        grid.addWidget(lbl_1, 13, 0)
        grid.addWidget(lbl_2, 14, 0)
        grid.addWidget(self.reysis, 13, 1)
        grid.addWidget(self.intervals, 14, 1)
        grid.addWidget(customer, 15, 0)
        grid.addWidget(self.customerEdit, 15, 1)
        grid.addWidget(shell, 16, 0)
        grid.addWidget(self.shellEdit, 16, 1)
        grid.addWidget(cat_well, 17, 0)
        grid.addWidget(self.cat_wellEdit, 17, 1)
        grid.addWidget(self.button, 18, 1)
        self.setLayout(grid)
        self.setWindowTitle('Геологический Журнал')

    @pyqtSlot()
    def extraction_1(self):
        wellEditValue = self.wellEdit.text()
        locatEditValue = self.locatEdit.text()
        altustEditValue = self.altustEdit.text()
        altrotEditValue = self.altrotEdit.text()
        drillstartEditValue = self.drillstartEdit.text()
        drillendEditValue = self.drillendEdit.text()
        depth_projEditValue = self.depth_projEdit.text()
        depth_factEditValue = self.depth_factEdit.text()
        horizont_projEditValue = self.horizont_projEdit.text()
        horizont_factEditValue = self.horizont_factEdit.text()
        well_placeEditValue = self.well_placeEdit.text()
        diametrEditValue = self.diametrEdit.text()
        reysisValue = self.reysis.text()
        intervalsValue = self.intervals.text()


class Window2(QWidget):
    """docstring for Window2"""

    def __init__(self, reysis):
        super(Window2, self).__init__()
        self.reysis = reysis
        self.rows = int(reysis)

        grid1 = QGridLayout(self)
        head_1 = QLabel('От')
        head_2 = QLabel('До')
        head_3 = QLabel('Вынос керна')
        grid1.addWidget(head_1, 0, 0)
        grid1.addWidget(head_2, 0, 1)
        grid1.addWidget(head_3, 0, 2)

        self.reys_begin = {}
        self.reys_ends = {}
        self.core = {}

        self.reys_beginValue = {}
        self.reys_endsValue = {}
        self.coreValue = {}

        self.reys_begin[0] = QLineEdit()
        self.reys_begin[0].setValidator(QDoubleValidator(1.0, 499.9, 2))
        grid1.addWidget(self.reys_begin[0])

        for i in range(self.rows):
            k_2 = QLineEdit()
            self.reys_ends[i] = k_2
            self.reys_ends[i].setValidator(QDoubleValidator(1.0, 499.9, 2))
            grid1.addWidget(self.reys_ends[i], i + 1, 1)

        for i in range(self.rows):
            k_3 = QLineEdit()
            self.core[i] = k_3
            self.core[i].setValidator(QDoubleValidator(0.1, 499.9, 2))
            grid1.addWidget(self.core[i], i + 1, 2)

        self.button1 = QPushButton(self)
        self.button1.setText('Ok')
        self.button1.clicked.connect(self.extraction_2)
        self.button1_1 = QPushButton(self)
        self.button1_1.setText('Назад')
        grid1.addWidget(self.button1_1, self.rows + 1, 1)
        grid1.addWidget(self.button1, self.rows + 1, 2)
        self.setWindowTitle('Интервалы бурения')

    @pyqtSlot()
    def extraction_2(self):
        for i in range(self.rows):
            self.reys_endsValue[i] = self.reys_ends[i].text()
            self.coreValue[i] = self.core[i].text()
            if i == 0:
                self.reys_beginValue[i] = self.reys_begin[i].text()
            else:
                self.reys_beginValue[i] = self.reys_ends[i-1].text()

class Window3(QWidget):
    def __init__(self, intervals, reysis, wellEdit, locatEdit, altustEdit, altrotEdit, drillstartEdit, drillendEdit,
                 depth_projEdit, depth_factEdit, horizont_projEdit,
                 horizont_factEdit, well_placeEdit, diametrEdit, authors_nameEdit, customerEdit, shellEdit,
                 cat_wellEdit, reys_beginValue, reys_endsValue, coreValue):
        super(Window3, self).__init__()
        self.intervals = intervals
        self.wrows = int(intervals)
        self.reysis = reysis
        self.rows = int(reysis)
        self.wellEdit = wellEdit
        self.locatEdit = locatEdit
        self.altustEdit = altustEdit
        self.altrotEdit = altrotEdit
        self.drillstartEdit = drillstartEdit
        self.drillendEdit = drillendEdit
        self.depth_projEdit = depth_projEdit
        self.depth_factEdit = depth_factEdit
        self.horizont_projEdit = horizont_projEdit
        self.horizont_factEdit = horizont_factEdit
        self.well_placeEdit = well_placeEdit
        self.diametrEdit = diametrEdit
        self.authors_nameEdit = authors_nameEdit
        self.customerEdit = customerEdit
        self.shellEdit = shellEdit
        self.cat_wellEdit = cat_wellEdit
        self.reys_beginValue = reys_beginValue
        self.reys_endsValue = reys_endsValue
        self.coreValue = coreValue

        grid2 = QGridLayout(self)
        head_4 = QLabel('От')
        head_5 = QLabel('До')
        head_6 = QLabel('Возраст')
        head_7 = QLabel('Описание')
        head_8 = QLabel('Примечания')
        grid2.addWidget(head_4, 0, 0)
        grid2.addWidget(head_5, 0, 1)
        grid2.addWidget(head_6, 0, 2)
        grid2.addWidget(head_7, 0, 3)
        grid2.addWidget(head_8, 0, 4)

        self.layer_begins = {}
        self.layer_ends = {}
        self.age_layer = {}
        self.discr = {}
        self.gnvp = {}
        self.layer_beginsValue = {}
        self.layer_endsValue = {}
        self.age_layerValue = {}
        self.discrValue = {}
        self.gnvpValue = {}

        sands = QLabel('Пройдено по песчаной пачке')
        self.sand_begin = QLineEdit()
        self.sand_end = QLineEdit()

        self.layer_begins[0] = QLineEdit()
        self.layer_begins[0].setValidator(QDoubleValidator(1.0, 499.9, 2))
        grid2.addWidget (self.layer_begins[0])

        for i in range(self.wrows):
            m_2 = QLineEdit()
            self.layer_ends[i] = m_2
            self.layer_ends[i].setValidator(QDoubleValidator(1.0, 499.9, 2))
            grid2.addWidget(self.layer_ends[i], i + 1, 1)

        for i in range(self.wrows):
            m_3 = QLineEdit()
            self.age_layer[i] = m_3
            grid2.addWidget(self.age_layer[i], i + 1, 2)

        for i in range(self.wrows):
            m_4 = QTextEdit()
            m_4.setMaximumHeight(50)
            self.discr[i] = m_4
            grid2.addWidget(self.discr[i], i + 1, 3)

        for i in range(self.wrows):
            m_5 = QLineEdit()
            self.gnvp[i] = m_5
            grid2.addWidget(self.gnvp[i], i + 1, 4)

        self.button2_1 = QPushButton(self)
        self.button2_1.setText('Назад')
        self.button2 = QPushButton(self)
        self.button2.setText('Ok')
        self.button2.clicked.connect(self.execute)

        grid2.addWidget(sands, self.wrows + 1, 0)
        grid2.addWidget(self.sand_begin, self.wrows + 1, 1)
        grid2.addWidget(self.sand_end, self.wrows + 1, 2)
        grid2.addWidget(self.button2_1, self.wrows + 2, 1)
        grid2.addWidget(self.button2, self.wrows + 2, 2)
        self.setWindowTitle('Описание керна')


    def execute(self):
        for i in range(self.wrows):
            self.layer_endsValue[i] = self.layer_ends[i].text()
            self.age_layerValue[i] = self.age_layer[i].text()
            self.discrValue[i] = self.discr[i].toPlainText()
            self.gnvpValue[i] = self.gnvp[i].text()
            if i == 0:
                self.layer_beginsValue[i] = self.layer_begins[i].text()
            else:
                self.layer_beginsValue[i] = self.layer_ends[i-1].text()

        self.document = Document()

        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        sections = self.document.sections
        sections.left_margin = Cm(1.0)
        sections.right_margin = Cm(1.0)

        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run('Геологический журнал').bold = True

        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run('ТГРУ')

        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(self.locatEdit)

        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run('СКВАЖИНА  №')
        paragraph.add_run(self.wellEdit)

        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run('Альтитуда устья: ')
        paragraph.add_run(self.altustEdit)

        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run('Альтитуда ротора: ')
        paragraph.add_run(self.altrotEdit)

        table = self.document.add_table(8, 2, "Table Grid")
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Скважина начата'
        heading_cells[1].text = self.drillstartEdit
        cell_1 = table.cell(1, 0)
        cell_1.text = 'Скважина закончана'
        cell_2 = table.cell(2, 0)
        cell_2.text = 'Глубина проектная'
        cell_3 = table.cell(3, 0)
        cell_3.text = 'Глубина фактическая'
        cell_4 = table.cell(4, 0)
        cell_4.text = 'Горизонт проектная'
        cell_5 = table.cell(5, 0)
        cell_5.text = 'Горизонт фактическая'
        cell_6 = table.cell(6, 0)
        cell_6.text = 'Местоположение скважины'
        cell_7 = table.cell(7, 0)
        cell_7.text = 'Диаметр ствола скважины'
        cell_2_1 = table.cell(1, 1)
        cell_2_1.text = self.drillendEdit
        cell_2_2 = table.cell(2, 1)
        cell_2_2.text = self.depth_projEdit
        cell_2_3 = table.cell(3, 1)
        cell_2_3.text = self.depth_factEdit
        cell_2_4 = table.cell(4, 1)
        cell_2_4.text = self.horizont_projEdit
        cell_2_5 = table.cell(5, 1)
        cell_2_5.text = self.horizont_factEdit
        cell_2_6 = table.cell(6, 1)
        cell_2_6.text = self.well_placeEdit
        xx_12 = '0,0' + ' - ' + self.depth_factEdit + '  -  ' + self.diametrEdit
        cell_2_7 = table.cell(7, 1)
        cell_2_7.text = xx_12

        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run('\n\n\n\n\n\n\n\n\nКонструкция скважины:')

        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        from docx.enum.table import WD_TABLE_ALIGNMENT
        table_1 = self.document.add_table(4, 5, "Table Grid")

        row = table_1.rows[0]
        a, b = row.cells[2:4]
        A = a.merge(b)
        q = ['Наименование колонны', 'Диаметр в мм.', '0', 'Интервал крепеления, м', 'Общая длинна труб, м']
        i = 0
        for i in range(0, 5):
            row.cells[i].text = q[i]
        cell_3 = table_1.cell(2, 0)
        cell_3.text = "Направление"
        cell_3_1 = table_1.cell(3, 0)
        cell_3_1.text = "Экспл. Колонна"
        cell_3_2 = table_1.cell(1, 2)
        cell_3_2.text = "От"
        cell_3_3 = table_1.cell(1, 3)
        cell_3_3.text = "До"
        row_1_1 = table_1.rows[2]
        row_1_2 = table_1.rows[3]
        for i in range(1, 5):
            row_1_1.cells[i].text = "-"
            row_1_2.cells[i].text = "-"
        self.document.add_page_break()
        # ------------------------------------------------------------------------------------------------------
        reys_lenth = []
        for i in range(self.rows):
            t_1 = float(self.reys_endsValue[i].replace(',', '.'))
            t_2 = float(self.reys_beginValue[i].replace(',', '.'))
            j_1 = t_1 - t_2
            j = round(j_1, 1)
            reys_lenth.append(j)
        paragraph = self.document.add_paragraph()
        table_4 = self.document.add_table(2, 5, "Table Grid")
        row = table_4.rows[0]
        a, b = row.cells[0:2]
        A = a.merge(b)
        a, b = row.cells[1:3]
        A = a.merge(b)
        a, b = row.cells[3:5]
        A = a.merge(b)
        cell_4_1 = table_4.cell(0, 2)
        cell_4_1.text = 'Пробурено'
        cell_4_2 = table_4.cell(0, 4)
        cell_4_2.text = 'Вынос керна'
        s_1 = ['От', 'До', 'Итого', 'М', '%']
        row_1 = table_4.rows[1]
        i = 0
        for i in range(0, 5):
            row_1.cells[i].text = s_1[i]
        i = 0
        for i in range(self.rows):
            row = table_4.add_row()
        column_1 = table_4.columns[0]
        for i in range(self.rows):
            k_4 = float((self.reys_beginValue[i]).replace(',', '.'))
            k_4 = round(k_4, 1)
            k_4 = str(k_4).replace('.', ',')
            column_1.cells[i + 2].text = k_4
        column_2 = table_4.columns[1]
        i = 0
        for i in range(self.rows):
            k_4 = round(float((self.reys_endsValue[i]).replace(',', '.')), 1)
            k_4 = str(k_4).replace('.', ',')
            column_2.cells[i + 2].text = k_4
        column_3 = table_4.columns[2]
        for i in range(self.rows):
            self.coreValue[i] = float(self.coreValue[i].replace(',', '.'))
        column_4 = table_4.columns[3]
        core_sum = 0
        core_sum = float(core_sum)
        column_5 = table_4.columns[4]
        W_1 = []
        for i in range(self.rows):
            self.coreValue[i] = float(self.coreValue[i])
            reys_lenth[i] = float(reys_lenth[i])
            core_sum = float(core_sum) + self.coreValue[i]
            W_1.append(round(self.coreValue[i] / reys_lenth[i] * 100, 1))
        for i in range(self.rows):
            W_1[i] = str(W_1[i])
            W_1[i] = W_1[i].replace('.', ',')
            column_5.cells[i + 2].text = W_1[i]
        for i in range(self.rows):
            self.coreValue[i] = str(self.coreValue[i])
            self.coreValue[i] = self.coreValue[i].replace('.', ',')
            column_4.cells[i + 2].text = self.coreValue[i]
        i = 0
        for i in range(self.rows):
            reys_lenth[i] = str(reys_lenth[i])
        for i in range(self.rows):
            reys_lenth[i] = reys_lenth[i].replace('.', ',')
            reys_lenth[i] = str(reys_lenth[i])
        i = 2
        for i in range(self.rows):
            column_3.cells[i + 2].text = reys_lenth[i]
        additional_row_0 = table_4.add_row()
        additional_row_1 = table_4.add_row()
        additional_row_2 = table_4.add_row()
        additional_row_1.cells[0].text = "по песч. пачке"
        additional_row_2.cells[0].text = "по прод. слою"

        sands_begin = self.sand_begin.text()
        sands_end = self.sand_end.text()
        sands_begin = float(sands_begin.replace(',', '.'))
        sands_end = float(sands_end.replace(',', '.'))
        sands_width = sands_end - sands_begin
        sands_width = round(sands_width, 1)
        additional_row_1.cells[1].text = str(sands_begin).replace('.', ',')
        additional_row_1.cells[2].text = str(sands_end).replace('.', ',')
        additional_row_1.cells[3].text = str(sands_width).replace('.', ',')
        a = additional_row_1.cells[4]
        b = additional_row_2.cells[4]
        A_11 = a.merge(b)
        drill_end = float(self.reys_endsValue[self.rows - 1].replace(',', '.'))
        drill_start = float(self.reys_beginValue[0].replace(',', '.'))
        drill_sum = (drill_end) - (drill_start)
        drill_sum = round(drill_sum, 1)
        core_percent = float(core_sum / drill_sum * 100)
        core_percent = round(core_percent, 1)
        core_percent_final = 'Средний выход ' + str(core_percent).replace('.', ',') + '%'
        A_11.text = core_percent_final
        self.document.add_page_break()
        # -------------------------------------------------описание керна----------------------------------------------------------------------------------------------------------------
        paragraph = self.document.add_paragraph()
        table_5 = self.document.add_table(3, 7, "Table Grid")
        row_2_0 = table_5.rows[0]
        a, b = row_2_0.cells[4:6]
        A = a.merge(b)
        row_2_1 = table_5.rows[1]
        a, b = row_2_1.cells[4:6]
        A = a.merge(b)
        row_2_2 = table_5.rows[2]
        a, b = row_2_2.cells[4:6]
        A = a.merge(b)
        row_2_0 = table_5.rows[0]
        a, b = row_2_0.cells[0:2]
        A = a.merge(b)
        q_1 = ['Глубина за-легания слоя, м', 'Мощ-ность слоя, м', 'Геоло-гичес. индекс', '0', 'Описание породы',
               'Водо-нефте-газо-проявл.']
        i = 0
        for i in range(0, 6):
            row_2_0.cells[i + 1].text = q_1[i]
        M_1 = []
        layers = []
        i = 0
        for i in range(self.wrows):
            stratum = table_5.add_row()
            strat = table_5.rows[3 + i]
            a, b = strat.cells[4:6]
            A = a.merge(b)
            strat.cells[4].width = Inches(3.7)
            la_1 = float(self.layer_beginsValue[i].replace(',', '.'))
            la_2 = float(self.layer_endsValue[i].replace(',', '.'))
            la_3 = la_2 - la_1
            la_3 = round(la_3, 1)
            la_4 = self.age_layerValue[i]
            la_5 = self.discrValue[i]
            la_6 = self.gnvpValue[i]
            layer = [la_1, la_2, la_3, la_4, la_5, la_6]
            layers.append(layer)
            strat.cells[0].text = str(la_1).replace('.', ',')
            strat.cells[1].text = str(la_2).replace('.', ',')
            strat.cells[2].text = str(la_3).replace('.', ',')
            strat.cells[3].text = str(la_4)
            strat.cells[4].text = str(la_5)
            strat.cells[6].text = str(la_6)
            if la_6 == 'Битум':
                M_1.append(la_1)
                M_1.append(la_2)
        row_3_0 = table_5.rows[1]
        row_3_1 = table_5.rows[2]
        print(M_1)
        s_2 = ['1', '2', '3', '4', '5', '5', '6']
        for i in (0, 1):
            row_3_0.cells[i].text = s_1[i]
        for i in range(0, 7):
            row_3_1.cells[i].text = s_2[i]
        if len(M_1) > 0:
            M_2 = []
            for x in range(1, (len(M_1)) - 1, 2):
                if M_1[x] == M_1[x + 1]:
                    M_2.append(x)
                    M_2.append(x + 1)
            for x in range((len(M_2))):
                M_1[M_2[x]] = 0

            for x in reversed(M_2):
                M_1.pop(x)
            if len(M_1) > 2:
                m_2 = M_1[0]
                m_1 = M_1[(len(M_1) - 1)]
                mmm = []
                for x in range(0, (len(M_1)) - 1, 2):
                    r = round(M_1[x + 1] - M_1[x], 1)
                    mmm.append(r)
                mmm = sum(mmm)
            else:
                m_1 = M_1[1]
                m_2 = M_1[0]
                mmm = round((m_1 - m_2), 1)

        else:
            m_2 = "-"
            m_1 = "-"
            mmm = ""

        additional_row_2.cells[1].text = str(m_2).replace('.', ',')
        additional_row_2.cells[2].text = str(m_1).replace('.', ',')
        additional_row_2.cells[3].text = str(mmm).replace('.', ',')

        authors = 'Описал геолог:              ' + self.authors_nameEdit
        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(authors)
        # АКТ СДАЧИ - ПРИЕМА КЕРНА ---------------------------------------------------------------------------------
        self.document.add_page_break()

        first_row_of_a_table = ['Заказчик:', self.customerEdit, 'Снаряд:', self.shellEdit]
        second_row_of_a_table = ['Категория скважины:', self.cat_wellEdit]

        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run('Акт\n')
        paragraph.add_run('сдачи - приема керна').bold = True

        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(
            'Мы, нижеподписавшиеся представители ТГРУ ПАО "ТАТНЕФТЬ"  геолог_______________________ с одной стороны и представители ______________________________________________ с другой стороны составили акт о том, что "___"______________2018 года первые сдали, а вторые приняли на хранение керн скважины')

        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run(
            '\nЦелевое назначение отбора керна: исследование коллекторских свойств нефтяной, водонефтяной, промытой зоны или опытного участка, уточнения геологического строения (необходимое подчеркнуть)')

        table_6 = self.document.add_table(2, 4)
        row_4 = table_6.rows[0]
        row_4_1 = table_6.rows[1]
        a, b = row_4_1.cells[1:3]
        A = a.merge(b)
        i = 0
        for i in range(0, 4):
            row_4.cells[i].text = first_row_of_a_table[i]
        for i in (2, 3):
            row_4_1.cells[i].text = second_row_of_a_table[i - 2]
        table_7_heading = ['интервал отбора, (м)', 'проходка, (м)', 'вынос, (м)', 'мощность нефт. керна (м)',
                           'возраст по интервалу', 'номер ящика']
        table_7 = self.document.add_table(self.rows + 1, 7, "Table Grid")
        table_7.autofit = True

        row_5_1 = table_7.rows[0]
        for i in range(0, 6):
            row_5_1.cells[i + 1].text = table_7_heading[i]
        a, b = row_5_1.cells[0:2]
        A = a.merge(b)

        column_2_1 = table_7.columns[0]
        column_2_2 = table_7.columns[1]
        column_2_3 = table_7.columns[2]
        #column_2_4 = table_7.columns [3]
        column_2_5 = table_7.columns[3]
        column_2_6 = table_7.columns[4]
        column_2_7 = table_7.columns[5]
        column_2_8 = table_7.columns[6]
        column_2_1.width = Inches(4.0)
        column_2_2.width = Inches(2.5)
        column_2_3.width = Inches(2.5)
        #column_2_4.width=Inches(2.5)
        column_2_5.width = Inches(2.5)
        column_2_6.width = Inches(2.5)
        column_2_7.width = Inches(2.5)
        column_2_7.width = Inches(2.5)

        for i in range(self.rows):
            k_5 = round(float((self.reys_beginValue[i]).replace(',', '.')), 1)
            k_5 = str(k_5).replace('.', ',')
            column_2_1.cells[i + 1].text = k_5
            column_2_1.cells[i + 1].width = Inches(4.5)
        for i in range(self.rows):
            k_6 = round(float((self.reys_endsValue[i]).replace(',', '.')), 1)
            k_6 = str(k_6).replace('.', ',')
            column_2_2.cells[i + 1].text = k_6
            column_2_2.cells[i + 1].width = Inches(4.5)
        for i in range(self.rows):
            column_2_3.cells[i + 1].text = reys_lenth[i]
        # for i in range (self.rows):
        # self.coreValue[i] = str(self.coreValue[i])
        # self.coreValue[i] = self.coreValue[i].replace ('.',',')
        # column_2_4.cells[i+1].text = self.coreValue[i]
        # self.coreValue[i] = self.coreValue[i].replace (',','.')
        # self.coreValue[i] = float(self.coreValue[i])
        O = []
        p = 0.0
        for i in range(0, self.rows):
            self.coreValue[i] = self.coreValue[i].replace(',', '.')
            self.coreValue[i] = round(float(self.coreValue[i]), 1)
            p = p + self.coreValue[i]
            p = round(p, 1)
            O.append(p)

        for i in range(self.rows):
            O[i] = str(O[i]).replace('.', ',')
            column_2_5.cells[i + 1].text = str(self.coreValue[i])
        for i in range(self.rows):
            O[i] = round(float(O[i].replace(',', '.')), 1)
            if (O[i] / 5) < 1:
                column_2_8.cells[i + 1].text = '1'
            elif (O[i] / 5) > 1:
                if (O[i] // 5) * 5 < O[i] and (O[i] // 5) * 5 > O[i] - self.coreValue[i]:
                    column_2_8.cells[i + 1].text = str(int(O[i] // 5)) + '-' + str(int((O[i] // 5) + 1))
                elif (O[i] // 5) * 5 < O[i] and (O[i] // 5) * 5 < O[i] - self.coreValue[i]:
                    column_2_8.cells[i + 1].text = str(int((O[i] // 5)) + 1)
                elif (O[i] // 5) * 5 == O[i]:
                    column_2_8.cells[i + 1].text = str(int((O[i] // 5)) + 1)
        column_2_6.cells[1].text = str(m_2).replace('.', ',') + '-' + str(m_1).replace('.', ',')
        column_2_6.cells[2].text = str(mmm).replace('.', ',')
        # -------------------------------------------------------------------------------------------------------------------------------------
        age_layer_list = []
        age_layer_outsole = []
        ages_outsoles = []
        reysis_ages = []

        for i in range(self.wrows):
            buf_2 = layers[i]
            k = buf_2[3]
            l = round(float(buf_2[1]), 1)
            age_layer_list.append(k)
            age_layer_outsole.append(l)
        new_x = [el for el, _ in groupby(age_layer_list)]
        print(new_x)
        print(age_layer_list)
        print(age_layer_outsole)
        for x in range((len(age_layer_list)) - 1):
            if age_layer_list[x] != age_layer_list[x + 1]:
                ages_outsoles.append(x)
            else:
                pass
        ages_outsoles.append((len(age_layer_list)) - 1)
        k_2 = round(float((self.layer_beginsValue[0]).replace(',', '.')), 1)
        age_layer_outsole.insert(0, k_2)
        print(age_layer_outsole)
        print(ages_outsoles)
        print(len(age_layer_outsole))
        for i in range(self.rows):
            self.reys_beginValue[i] = round(float(self.reys_beginValue[i].replace(',', '.')), 1)
            self.reys_endsValue[i] = round(float(self.reys_endsValue[i].replace(',', '.')), 1)

        for i in range(self.rows):
            for j in range(1, (len(age_layer_outsole))):
                if self.reys_beginValue[i] >= age_layer_outsole[j - 1] and self.reys_endsValue[i] <= age_layer_outsole[
                    j]:
                    reysis_ages.append(age_layer_list[j - 1])
                if self.reys_beginValue[i] >= age_layer_outsole[j - 1] and self.reys_endsValue[i] > age_layer_outsole[
                    j] and self.reys_beginValue[i] < age_layer_outsole[j]:
                    if age_layer_list[j - 1] == age_layer_list[j]:
                        reysis_ages.append(age_layer_list[j - 1])
                    else:
                        reysis_ages.append(age_layer_list[j - 1] + '+' + age_layer_list[j])
                else:
                    pass
        print(reysis_ages)
        for i in range(self.rows):
            column_2_7.cells[i + 1].text = reysis_ages[i]

class Window4(QWidget):
    def __init__(self, document):
        super(Window4, self).__init__()
        self.document = document
        self.title = 'Сохранение'
        self.left = 10
        self.top = 10
        self.width = 640
        self.height = 480
        self.initUI()

    def initUI(self):
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)
        self.saveFileDialog()
        self.show()

    def saveFileDialog(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self,"Выберете путь сохранения","","Word Files (*.docx)", options=options)
        if fileName:
            self.document.save(fileName+ '.docx')

#######################################################################################################################################################
class MainWindow(QMainWindow):
    """docstring for MainWindow"""

    def __init__(self, ):
        super(MainWindow, self).__init__()

    def show_window_1(self):
        self.w1 = Window1()
        self.w1.button.clicked.connect(self.show_window_2)
        self.w1.button.clicked.connect(self.w1.close)
        self.w1.show()

    def show_window_2(self):
        self.w2 = Window2(self.w1.reysis.text())
        self.w2.button1.clicked.connect(self.show_window_3)
        self.w2.button1.clicked.connect(self.w2.close)
        self.w2.button1_1.clicked.connect(self.show_window_1)
        self.w2.show()

    def show_window_3(self):
        self.w3 = Window3(self.w1.intervals.text(), self.w1.reysis.text(), self.w1.wellEdit.text(),
                          self.w1.locatEdit.text(), self.w1.altustEdit.text(),
                          self.w1.altrotEdit.text(), self.w1.drillstartEdit.text(), self.w1.drillendEdit.text(),
                          self.w1.depth_projEdit.text(),
                          self.w1.depth_factEdit.text(), self.w1.horizont_projEdit.text(),
                          self.w1.horizont_factEdit.text(), self.w1.well_placeEdit.text(),
                          self.w1.diametrEdit.text(), self.w1.authors_nameEdit.text(), self.w1.customerEdit.text(),
                          self.w1.shellEdit.text(), self.w1.cat_wellEdit.text(),
                          self.w2.reys_beginValue, self.w2.reys_endsValue, self.w2.coreValue)
        self.w3.button2_1.clicked.connect(self.show_window_2)
        self.w3.button2_1.clicked.connect(self.w3.close)
        self.w3.button2.clicked.connect(self.show_window_4)
        self.w3.button2.clicked.connect(self.w3.close)
        self.w3.show()
    def show_window_4(self):
        self.w4 = Window4(self.w3.document)
        self.w4.close()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    w = MainWindow()
    w.show_window_1()
    sys.exit(app.exec_())